# -*- coding: utf-8 -*-
"""
读取 Word 文档 (.doc/.docx) 的脚本 - 保留文档布局
"""

import os
import sys


def get_heading_level(paragraph):
    """获取标题级别"""
    try:
        style_name = paragraph.style.name.lower()
        if 'heading 1' in style_name or '标题 1' in style_name:
            return 1
        elif 'heading 2' in style_name or '标题 2' in style_name:
            return 2
        elif 'heading 3' in style_name or '标题 3' in style_name:
            return 3
        elif 'heading' in style_name or '标题' in style_name:
            return 1
    except:
        pass
    return 0


def read_docx_with_format(file_path):
    """读取 .docx 文件内容，保留格式信息"""
    try:
        from docx import Document
    except ImportError:
        print("请先安装 python-docx: pip install python-docx")
        return None
    
    try:
        doc = Document(file_path)
        result = []
        
        # 读取段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                result.append('')  # 保留空行
                continue
            
            # 判断是否为标题
            heading_level = get_heading_level(para)
            
            if heading_level > 0:
                # 标题格式化
                prefix = '#' * heading_level + ' '
                result.append(f"\n{prefix}{text}\n")
            else:
                # 普通段落
                # 根据缩进添加空格
                indent = ''
                try:
                    if para.paragraph_format.first_line_indent:
                        indent = '    '  # 首行缩进4个空格
                except:
                    pass
                result.append(f"{indent}{text}")
        
        # 读取表格
        if doc.tables:
            result.append("\n" + "="*60)
            result.append("【表格内容】")
            result.append("="*60)
            
            for table_idx, table in enumerate(doc.tables, 1):
                result.append(f"\n--- 表格 {table_idx} ---")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip().replace('\n', ' ')
                        row_text.append(cell_text)
                    result.append(' | '.join(row_text))
                result.append('')
        
        return '\n'.join(result)
        
    except Exception as e:
        print(f"读取文件时出错: {e}")
        return None


def read_doc(file_path):
    """读取 .doc 文件内容 (使用 antiword 或 pywin32)"""
    if sys.platform == 'win32':
        word = None
        doc = None
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.visible = False
            doc = word.Documents.Open(os.path.abspath(file_path))
            
            # 尝试获取带格式的文本
            result = []
            
            # 遍历段落获取格式信息
            for para in doc.Paragraphs:
                try:
                    text = para.Range.Text.strip()
                    if not text:
                        result.append('')
                        continue
                    
                    # 获取样式信息
                    style_name = ""
                    try:
                        style_name = para.Style.NameLocal
                    except:
                        pass
                    
                    # 根据样式判断标题级别
                    is_heading = False
                    level = 0
                    if style_name:
                        style_lower = style_name.lower()
                        if '标题' in style_name or 'heading' in style_lower:
                            is_heading = True
                            if '1' in style_name:
                                level = 1
                            elif '2' in style_name:
                                level = 2
                            elif '3' in style_name:
                                level = 3
                            else:
                                level = 1
                    
                    if is_heading:
                        prefix = '#' * level + ' '
                        result.append(f"\n{prefix}{text}\n")
                    else:
                        # 检查缩进
                        indent = ''
                        try:
                            if para.FirstLineIndent and para.FirstLineIndent > 0:
                                indent = '    '
                        except:
                            pass
                        result.append(f"{indent}{text}")
                except Exception as para_error:
                    # 如果某个段落处理失败，跳过该段落
                    continue
            
            # 读取表格
            try:
                if doc.Tables.Count > 0:
                    result.append("\n" + "="*60)
                    result.append("【表格内容】")
                    result.append("="*60)
                    
                    for i in range(1, doc.Tables.Count + 1):
                        try:
                            table = doc.Tables(i)
                            result.append(f"\n--- 表格 {i} ---")
                            for row in range(1, table.Rows.Count + 1):
                                row_text = []
                                for col in range(1, table.Columns.Count + 1):
                                    try:
                                        cell_text = table.Cell(row, col).Range.Text.strip()
                                        cell_text = cell_text.replace('\r', '').replace('\x07', '')
                                        row_text.append(cell_text)
                                    except:
                                        row_text.append('')
                                result.append(' | '.join(row_text))
                            result.append('')
                        except:
                            continue
            except:
                pass  # 表格读取失败则跳过
            
            return '\n'.join(result)
            
        except ImportError:
            print("Windows 平台需要安装 pywin32: pip install pywin32")
        except Exception as e:
            print(f"使用 Word COM 接口读取时出错: {e}")
        finally:
            # 确保关闭文档和Word应用
            try:
                if doc:
                    doc.Close(SaveChanges=False)
            except:
                pass
            try:
                if word:
                    word.Quit()
            except:
                pass
    
    try:
        import subprocess
        result = subprocess.run(['antiword', file_path], capture_output=True, text=True)
        if result.returncode == 0:
            return result.stdout
        else:
            print(f"antiword 读取失败: {result.stderr}")
    except FileNotFoundError:
        print("未找到 antiword，请安装: sudo apt-get install antiword (Linux)")
    except Exception as e:
        print(f"使用 antiword 读取时出错: {e}")
    
    return None


def read_word_document(file_path):
    """根据文件扩展名自动选择读取方式"""
    if not os.path.exists(file_path):
        print(f"文件不存在: {file_path}")
        return None
    
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.docx':
        return read_docx_with_format(file_path)
    elif ext == '.doc':
        return read_doc(file_path)
    else:
        print(f"不支持的文件格式: {ext}")
        return None


def main():
    # 示例用法
    file_path = r"e:\毕设\新建文件夹\论文文档\任务书.doc"  # 修改为你的文件路径
    
    print(f"正在读取文件: {file_path}")
    print("="*60)
    content = read_word_document(file_path)
    
    if content:
        print(content)
        print("="*60)
    else:
        print("无法读取文件内容")


if __name__ == "__main__":
    main()
